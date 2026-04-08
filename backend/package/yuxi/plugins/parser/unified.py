"""Unified parser module for markdown conversion."""

from __future__ import annotations

import asyncio
import base64
import os
import re
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import aiofiles
from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter
from langchain_community.document_loaders import PyPDFLoader
from markdownify import markdownify as md_convert

from yuxi.plugins.parser.zip_utils import process_zip_file as _process_zip_file
from yuxi.storage.minio import get_minio_client
from yuxi.utils import logger

SUPPORTED_FILE_EXTENSIONS: tuple[str, ...] = (
    ".txt",
    ".md",
    ".docx",
    ".html",
    ".htm",
    ".json",
    ".csv",
    ".xls",
    ".xlsx",
    ".pdf",
    ".pptx",
    ".jpg",
    ".jpeg",
    ".png",
    ".bmp",
    ".tiff",
    ".tif",
    ".zip",
)


def is_supported_file_extension(file_name: str | os.PathLike[str]) -> bool:
    """Check whether the given file path has a supported extension."""
    return Path(file_name).suffix.lower() in SUPPORTED_FILE_EXTENSIONS


@dataclass(slots=True)
class MarkdownParseResult:
    """统一的 Markdown 解析结果。"""

    markdown: str
    file_ext: str | None = None
    artifacts: dict[str, Any] = field(default_factory=dict)


_docling_converter: DocumentConverter | None = None


def _get_docling_converter() -> DocumentConverter:
    """获取 Docling 文档转换器单例。"""
    global _docling_converter
    if _docling_converter is None:
        _docling_converter = DocumentConverter(
            format_options={
                InputFormat.DOCX: None,
                InputFormat.XLSX: None,
                InputFormat.PPTX: None,
            }
        )
    return _docling_converter


def _resolve_image_storage_params(params: dict | None) -> tuple[str, str]:
    params = params or {}

    image_bucket = params.get("image_bucket") or "public"
    image_prefix = params.get("image_prefix")
    if image_prefix:
        normalized_prefix = str(image_prefix).strip("/")
        if normalized_prefix:
            return image_bucket, normalized_prefix

    return image_bucket, "unknown/kb-images"


def _upload_image_to_minio(image_data: bytes, filename: str, bucket_name: str, object_prefix: str) -> str:
    """上传图片到 MinIO，返回 URL。"""
    minio_client = get_minio_client()
    minio_client.ensure_bucket_exists(bucket_name)

    normalized_prefix = object_prefix.strip("/") or "unknown/kb-images"
    timestamp = int(time.time() * 1000000)
    object_name = f"{normalized_prefix}/{timestamp}_{Path(filename).name}"

    result = minio_client.upload_file(
        bucket_name=bucket_name,
        object_name=object_name,
        data=image_data,
    )
    return result.url


def _parse_data_uri(data_uri: str) -> tuple[bytes, str]:
    """解析 data URI，返回 (image_data, mime_type)。"""
    header, base64_data = data_uri.split(",", 1)
    mime_type = header.split(":")[1].split(";")[0]
    image_data = base64.b64decode(base64_data)
    return image_data, mime_type


def _convert_with_docling(file_path: Path, params: dict | None = None) -> str:
    """使用 Docling 将 docx/xlsx/pptx 转换为 Markdown。"""
    params = params or {}
    image_bucket, image_prefix = _resolve_image_storage_params(params)

    converter = _get_docling_converter()
    result = converter.convert(file_path)

    if result.status.name != "SUCCESS":
        raise RuntimeError(f"Docling 转换失败: {result.status}")

    doc = result.document

    if hasattr(doc, "pictures") and doc.pictures:
        image_refs: list[tuple[str, bytes]] = []

        for pic in doc.pictures:
            if hasattr(pic, "image") and hasattr(pic.image, "uri"):
                uri = str(pic.image.uri)
                if uri.startswith("data:"):
                    image_data, mime_type = _parse_data_uri(uri)
                    timestamp = int(time.time() * 1000000)
                    filename = f"image_{timestamp}.{mime_type.split('/')[-1]}"
                    image_refs.append((filename, image_data))

        image_urls: list[str] = []
        for filename, image_data in image_refs:
            try:
                url = _upload_image_to_minio(image_data, filename, image_bucket, image_prefix)
                image_urls.append(f"![{filename}]({url})")
            except Exception as e:  # noqa: BLE001
                logger.error(f"上传图片失败 {filename}: {e}")
                image_urls.append(f"[图片: {filename}]")

        markdown = doc.export_to_markdown()

        for url in image_urls:
            markdown = re.sub(r"<!--\s*image\s*-->", url, markdown, count=1)

        return markdown

    return doc.export_to_markdown()


def _convert_docx_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 解析 DOCX（Docling 失败时兜底）。"""
    from docx import Document

    document = Document(str(file_path))
    blocks: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        header = rows[0]
        blocks.append(f"| {' | '.join(header)} |")
        blocks.append(f"| {' | '.join(['---'] * len(header))} |")

        for row in rows[1:]:
            normalized_row = row + [""] * (len(header) - len(row))
            blocks.append(f"| {' | '.join(normalized_row[: len(header)])} |")

        blocks.append("")

    return "\n\n".join(blocks).strip()


def pdfreader(file_path, params=None):
    """读取 PDF 文件并返回 text 文本。"""
    if isinstance(file_path, str):
        file_path = Path(file_path)

    assert file_path.exists(), "File not found"
    assert file_path.suffix.lower() == ".pdf", "File format not supported"

    loader = PyPDFLoader(str(file_path))
    docs = loader.load()
    text = "\n\n".join([d.page_content for d in docs])
    return text


def parse_pdf(file, params=None):
    """解析 PDF 文件，支持多种 OCR 方式。"""
    from yuxi.plugins.parser.base import DocumentProcessorException
    from yuxi.plugins.parser.factory import DocumentProcessorFactory

    params = params or {}
    opt_ocr = params.get("enable_ocr", "disable")

    if opt_ocr == "disable":
        return pdfreader(file, params=params)

    image_bucket, image_prefix = _resolve_image_storage_params(params)
    params.setdefault("image_bucket", image_bucket)
    params.setdefault("image_prefix", image_prefix)

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, params)
    except DocumentProcessorException as e:
        logger.error(f"文档处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"PDF 解析失败: {str(e)}")
        raise DocumentProcessorException(f"PDF解析失败: {str(e)}", opt_ocr, "parsing_failed")


def parse_image(file, params=None):
    """解析图像文件，支持多种 OCR 方式。"""
    from yuxi.plugins.parser.base import DocumentProcessorException
    from yuxi.plugins.parser.factory import DocumentProcessorFactory

    params = params or {}
    opt_ocr = params.get("enable_ocr", "disable")

    if opt_ocr == "disable":
        raise ValueError(
            "图像文件必须启用OCR才能提取文本内容。"
            "请选择OCR方式 (rapid_ocr/mineru_ocr/mineru_official/pp_structure_v3_ocr) 或移除该文件。"
        )

    image_bucket, image_prefix = _resolve_image_storage_params(params)
    params.setdefault("image_bucket", image_bucket)
    params.setdefault("image_prefix", image_prefix)

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, params)
    except DocumentProcessorException as e:
        logger.error(f"图像处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:  # noqa: BLE001
        logger.error(f"图像解析失败: {str(e)}")
        raise DocumentProcessorException(f"图像解析失败: {str(e)}", opt_ocr, "parsing_failed")


async def parse_pdf_async(file, params=None):
    return await asyncio.to_thread(parse_pdf, file, params=params)


async def parse_image_async(file, params=None):
    return await asyncio.to_thread(parse_image, file, params=params)


async def _process_file_to_markdown_core(
    file_path: str, params: dict | None = None
) -> tuple[str, str | None, dict[str, Any]]:
    """将不同类型的文件转换为 markdown，支持本地文件和 MinIO 文件。"""
    from yuxi.knowledge.utils.kb_utils import is_minio_url, parse_minio_url
    from yuxi.storage.minio.client import get_minio_client

    if is_minio_url(file_path):
        logger.debug(f"Downloading file from MinIO: {file_path}")

        if "?" in file_path:
            file_path_clean = file_path.split("?")[0]
        else:
            file_path_clean = file_path

        original_filename = file_path_clean.split("/")[-1]

        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_filename).suffix) as temp_file:
            temp_path = temp_file.name

        try:
            bucket_name, object_name = parse_minio_url(file_path)
            minio_client = get_minio_client()
            file_content = await minio_client.adownload_file(bucket_name, object_name)

            async with aiofiles.open(temp_path, "wb") as f:
                await f.write(file_content)

            logger.debug(f"File downloaded to temp path: {temp_path}")
            actual_file_path = temp_path

        except Exception as e:  # noqa: BLE001
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            logger.error(f"Failed to download file from MinIO: {e}")
            raise ValueError(f"无法从MinIO下载文件: {e}")
    else:
        actual_file_path = file_path

    file_ext: str | None = None
    artifacts: dict[str, Any] = {}

    try:
        file_path_obj = Path(actual_file_path)
        file_ext = file_path_obj.suffix.lower()

        if file_ext == ".pdf":
            text = await parse_pdf_async(str(file_path_obj), params=params)
            result = f"{text}"

        elif file_ext in [".txt", ".md"]:
            with open(file_path_obj, encoding="utf-8") as f:
                content = f.read()
            result = f"{content}"

        elif file_ext == ".docx":
            try:
                result = _convert_with_docling(file_path_obj, params=params)
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Docling 解析 DOCX 失败，回退到 python-docx: {file_path_obj.name}, {e}")
                result = _convert_docx_with_python_docx(file_path_obj)

        elif file_ext == ".pptx":
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".doc":
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader

            loader = UnstructuredWordDocumentLoader(str(file_path_obj))
            docs = loader.load()
            result = "\n".join(doc.page_content for doc in docs).strip()

        elif file_ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"]:
            text = await parse_image_async(str(file_path_obj), params=params)
            result = f"{text}"

        elif file_ext in [".html", ".htm"]:
            with open(file_path_obj, encoding="utf-8") as f:
                content = f.read()
            text = md_convert(content, heading_style="ATX")
            result = f"{text}"

        elif file_ext == ".csv":
            import pandas as pd

            df = pd.read_csv(file_path_obj)
            markdown_content = ""

            for _, row in df.iterrows():
                row_df = pd.DataFrame([row], columns=df.columns)
                markdown_table = row_df.to_markdown(index=False)
                markdown_content += f"{markdown_table}\n\n"

            result = markdown_content.strip()

        elif file_ext in [".xls", ".xlsx"]:
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".json":
            import json

            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            data = json.loads(content)
            json_str = json.dumps(data, ensure_ascii=False, indent=2)
            result = f"```json\n{json_str}\n```"

        elif file_ext == ".zip":
            image_bucket, image_prefix = _resolve_image_storage_params(params)
            zip_result = await _process_zip_file(
                str(file_path_obj),
                image_bucket=image_bucket,
                image_prefix=image_prefix,
            )

            artifacts = {
                "zip_images_info": zip_result["images_info"],
                "zip_content_hash": zip_result["content_hash"],
                "zip_image_bucket": image_bucket,
                "zip_image_prefix": image_prefix,
            }

            result = zip_result["markdown_content"]

        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception:
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as cleanup_e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {cleanup_e}")
        raise

    finally:
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as e:  # noqa: BLE001
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {e}")

    return result, file_ext, artifacts


async def parse_source_to_markdown(source: str, params: dict | None = None) -> MarkdownParseResult:
    """统一入口: 将文件解析为 Markdown（URL 解析已废弃）。"""
    markdown, file_ext, artifacts = await _process_file_to_markdown_core(source, params=params)
    return MarkdownParseResult(
        markdown=markdown,
        file_ext=file_ext,
        artifacts=artifacts,
    )


class Parser:
    """Lightweight facade for converting file sources to markdown."""

    @staticmethod
    async def aparse(source: str, params: dict | None = None) -> str:
        """Asynchronously parse source content and return markdown text."""
        parsed = await parse_source_to_markdown(source=source, params=params)
        return parsed.markdown

    @classmethod
    def parse(cls, source: str, params: dict | None = None) -> str:
        """Synchronously parse source content and return markdown text."""
        try:
            asyncio.get_running_loop()
        except RuntimeError:
            return asyncio.run(cls.aparse(source=source, params=params))

        raise RuntimeError("当前处于异步上下文，请使用 `await Parser.aparse(...)`")
