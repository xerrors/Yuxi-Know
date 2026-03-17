import asyncio
import base64
import os
import re
import time
from pathlib import Path

import aiofiles
from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter
from langchain_community.document_loaders import (
    CSVLoader,
    JSONLoader,
    PyPDFLoader,
    TextLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from markdownify import markdownify as md_convert

from yuxi.plugins.parser.zip_utils import process_zip_file as _process_zip_file
from yuxi.storage.minio import get_minio_client
from yuxi.utils import logger

SUPPORTED_FILE_EXTENSIONS: tuple[str, ...] = (
    ".txt",
    ".md",
    ".docx",
    ".html",
    ".htm",
    ".json",
    ".csv",
    ".xls",
    ".xlsx",
    ".pdf",
    ".pptx",
    ".jpg",
    ".jpeg",
    ".png",
    ".bmp",
    ".tiff",
    ".tif",
    ".zip",
)


def is_supported_file_extension(file_name: str | os.PathLike[str]) -> bool:
    """Check whether the given file path has a supported extension."""
    return Path(file_name).suffix.lower() in SUPPORTED_FILE_EXTENSIONS


# Docling 文档转换器（单例模式）
_docling_converter: DocumentConverter | None = None


def _get_docling_converter() -> DocumentConverter:
    """获取 Docling 文档转换器单例"""
    global _docling_converter
    if _docling_converter is None:
        _docling_converter = DocumentConverter(
            format_options={
                InputFormat.DOCX: None,
                InputFormat.XLSX: None,
                InputFormat.PPTX: None,
            }
        )
    return _docling_converter


def _resolve_image_storage_params(params: dict | None) -> tuple[str, str]:
    params = params or {}

    image_bucket = params.get("image_bucket") or "public"
    image_prefix = params.get("image_prefix")
    if image_prefix:
        normalized_prefix = str(image_prefix).strip("/")
        if normalized_prefix:
            return image_bucket, normalized_prefix

    db_id = params.get("db_id")
    if db_id:
        return image_bucket, f"{db_id}/kb-images"

    return image_bucket, "unknown/kb-images"


def _upload_image_to_minio(image_data: bytes, filename: str, bucket_name: str, object_prefix: str) -> str:
    """上传图片到 MinIO，返回 URL"""
    minio_client = get_minio_client()
    minio_client.ensure_bucket_exists(bucket_name)

    normalized_prefix = object_prefix.strip("/") or "unknown/kb-images"
    timestamp = int(time.time() * 1000000)
    suffix = Path(filename).suffix.lower()
    object_name = f"{normalized_prefix}/{timestamp}_{Path(filename).name}"
    content_type_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }
    content_type = content_type_map.get(suffix, "image/jpeg")

    result = minio_client.upload_file(
        bucket_name=bucket_name,
        object_name=object_name,
        data=image_data,
        content_type=content_type,
    )
    return result.url


def _parse_data_uri(data_uri: str) -> tuple[bytes, str]:
    """解析 data URI，返回 (image_data, mime_type)"""
    header, base64_data = data_uri.split(",", 1)
    mime_type = header.split(":")[1].split(";")[0]
    image_data = base64.b64decode(base64_data)
    return image_data, mime_type


def _convert_with_docling(file_path: Path, params: dict | None = None) -> str:
    """
    使用 Docling 将 docx/xlsx/pptx 转换为 Markdown

    Args:
        file_path: 文件路径
        params: 参数，可包含 image_bucket/image_prefix

    Returns:
        Markdown 字符串
    """
    params = params or {}
    image_bucket, image_prefix = _resolve_image_storage_params(params)

    converter = _get_docling_converter()
    result = converter.convert(file_path)

    if result.status.name != "SUCCESS":
        raise RuntimeError(f"Docling 转换失败: {result.status}")

    doc = result.document

    # 提取图片并上传到 MinIO
    if hasattr(doc, "pictures") and doc.pictures:
        image_refs: list[tuple[str, bytes]] = []

        for pic in doc.pictures:
            if hasattr(pic, "image") and hasattr(pic.image, "uri"):
                uri = str(pic.image.uri)
                if uri.startswith("data:"):
                    image_data, mime_type = _parse_data_uri(uri)
                    timestamp = int(time.time() * 1000000)  # 微秒级时间戳
                    filename = f"image_{timestamp}.{mime_type.split('/')[-1]}"
                    image_refs.append((filename, image_data))

        # 上传图片并收集 URL
        image_urls: list[str] = []
        for filename, image_data in image_refs:
            try:
                url = _upload_image_to_minio(image_data, filename, image_bucket, image_prefix)
                image_urls.append(f"![{filename}]({url})")
            except Exception as e:
                logger.error(f"上传图片失败 {filename}: {e}")
                image_urls.append(f"[图片: {filename}]")

        # 导出 Markdown
        markdown = doc.export_to_markdown()

        # 替换 <!-- image --> 占位符为图片 URL
        # Docling 使用 <!-- image --> 作为占位符
        for url in reversed(image_urls):
            markdown = re.sub(r"<!--\s*image\s*-->", url, markdown, count=1)

        return markdown

    # 无图片时直接导出
    return doc.export_to_markdown()


def _convert_docx_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 解析 DOCX（Docling 失败时兜底）"""
    from docx import Document

    document = Document(str(file_path))
    blocks: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        header = rows[0]
        blocks.append(f"| {' | '.join(header)} |")
        blocks.append(f"| {' | '.join(['---'] * len(header))} |")

        for row in rows[1:]:
            normalized_row = row + [""] * (len(header) - len(row))
            blocks.append(f"| {' | '.join(normalized_row[: len(header)])} |")

        blocks.append("")

    return "\n\n".join(blocks).strip()


def chunk_with_parser(file_path, params=None):
    """
    使用文件解析器将文件切分成固定大小的块

    Args:
        file_path: 文件路径
        params: 参数
    """
    params = params or {}
    chunk_size = int(params.get("chunk_size", 500))
    chunk_overlap = int(params.get("chunk_overlap", 100))

    file_type = Path(file_path).suffix.lower()

    # 选择合适的加载器
    if file_type in [".txt"]:
        loader = TextLoader(file_path)

    elif file_type in [".md"]:
        loader = UnstructuredMarkdownLoader(file_path)

    elif file_type in [".docx", ".doc"]:
        loader = UnstructuredWordDocumentLoader(file_path)

    elif file_type in [".html", ".htm"]:
        loader = UnstructuredHTMLLoader(file_path)

    elif file_type in [".json"]:
        loader = JSONLoader(file_path, jq_schema=".")

    elif file_type in [".csv"]:
        loader = CSVLoader(file_path)

    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

    # 加载文档
    docs = loader.load()

    # 创建文本分割器
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    # 分割文档
    nodes = text_splitter.split_documents(docs)

    # 添加序号信息到metadata
    for i, node in enumerate(nodes):
        if node.metadata is None:
            node.metadata = {}
        node.metadata["chunk_idx"] = i

    return nodes


def chunk_text(text, params=None):
    """
    将文本切分成固定大小的块
    """
    params = params or {}
    chunk_size = int(params.get("chunk_size", 500))
    chunk_overlap = int(params.get("chunk_overlap", 100))

    # 创建文本分割器
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap, separators=["\n\n", "\n", ".", " ", ""]
    )

    # 分割文档
    nodes = text_splitter.split_text(text)

    # 添加序号信息到metadata
    nodes = [{"text": node, "metadata": {"chunk_idx": i}} for i, node in enumerate(nodes)]
    return nodes


def chunk(text_or_path, params=None):
    raise NotImplementedError("chunk is deprecated, use chunk_with_parser or chunk_text instead")


def pdfreader(file_path, params=None):
    """读取PDF文件并返回text文本"""
    if isinstance(file_path, str):
        file_path = Path(file_path)

    assert file_path.exists(), "File not found"
    assert file_path.suffix.lower() == ".pdf", "File format not supported"

    # 使用LangChain的PDF加载器
    loader = PyPDFLoader(str(file_path))
    docs = loader.load()

    # 简单的拼接起来之后返回纯文本
    text = "\n\n".join([d.page_content for d in docs])
    return text


def plainreader(file_path):
    """读取普通文本文件并返回text文本"""
    assert os.path.exists(file_path), "File not found"

    # 使用LangChain的文本加载器
    loader = TextLoader(str(file_path))
    docs = loader.load()
    text = "\n\n".join([d.page_content for d in docs])
    return text


def parse_pdf(file, params=None):
    """
    解析PDF文件，支持多种OCR方式

    Args:
        file: PDF文件路径
        params: 参数字典，包含enable_ocr设置

    Returns:
        str: 解析得到的文本

    Raises:
        DocumentProcessorException: 处理失败时抛出
    """
    from yuxi.plugins.parser.base import DocumentProcessorException
    from yuxi.plugins.parser.factory import DocumentProcessorFactory

    params = params or {}
    opt_ocr = params.get("enable_ocr", "disable")

    if opt_ocr == "disable":
        return pdfreader(file, params=params)

    image_bucket, image_prefix = _resolve_image_storage_params(params)
    params.setdefault("image_bucket", image_bucket)
    params.setdefault("image_prefix", image_prefix)

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, params)

    except DocumentProcessorException as e:
        logger.error(f"文档处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:
        logger.error(f"PDF 解析失败: {str(e)}")
        raise DocumentProcessorException(f"PDF解析失败: {str(e)}", opt_ocr, "parsing_failed")


def parse_image(file, params=None):
    """
    解析图像文件，支持多种OCR方式

    Args:
        file: 图像文件路径
        params: 参数字典，包含enable_ocr设置

    Returns:
        str: 解析得到的文本

    Raises:
        DocumentProcessorException: 处理失败时抛出
        ValueError: 图像文件禁用OCR时抛出
    """
    from yuxi.plugins.parser.base import DocumentProcessorException
    from yuxi.plugins.parser.factory import DocumentProcessorFactory

    params = params or {}
    opt_ocr = params.get("enable_ocr", "disable")

    # 图像文件必须使用 OCR,不能禁用
    if opt_ocr == "disable":
        raise ValueError(
            "图像文件必须启用OCR才能提取文本内容。"
            "请选择OCR方式 (rapid_ocr/mineru_ocr/mineru_official/pp_structure_v3_ocr) 或移除该文件。"
        )

    image_bucket, image_prefix = _resolve_image_storage_params(params)
    params.setdefault("image_bucket", image_bucket)
    params.setdefault("image_prefix", image_prefix)

    try:
        return DocumentProcessorFactory.process_file(opt_ocr, file, params)

    except DocumentProcessorException as e:
        logger.error(f"图像处理失败: {e.service_name} - {str(e)}")
        raise
    except Exception as e:
        logger.error(f"图像解析失败: {str(e)}")
        raise DocumentProcessorException(f"图像解析失败: {str(e)}", opt_ocr, "parsing_failed")


async def parse_pdf_async(file, params=None):
    return await asyncio.to_thread(parse_pdf, file, params=params)


async def parse_image_async(file, params=None):
    return await asyncio.to_thread(parse_image, file, params=params)


async def process_file_to_markdown(file_path: str, params: dict | None = None) -> str:
    """
    将不同类型的文件转换为markdown格式 - 支持本地文件和MinIO文件

    Args:
        file_path: 文件路径或MinIO URL
        params: 处理参数，对于ZIP文件需要包含 db_id

    Returns:
        markdown格式内容

    Note:
        对于ZIP文件，会在params中保存处理结果供调用方使用：
        - params['_zip_images_info']: 图片信息列表
        - params['_zip_content_hash']: 内容哈希值
    """
    import os
    import tempfile

    # 检测是否是MinIO URL
    from yuxi.knowledge.utils.kb_utils import is_minio_url

    if is_minio_url(file_path):
        # 从MinIO下载文件到临时位置
        logger.debug(f"Downloading file from MinIO: {file_path}")

        # 从MinIO URL中提取文件名
        if "?" in file_path:
            file_path_clean = file_path.split("?")[0]
        else:
            file_path_clean = file_path

        original_filename = file_path_clean.split("/")[-1]

        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_filename).suffix) as temp_file:
            temp_path = temp_file.name

        try:
            # 使用通用函数解析MinIO URL并下载文件
            from yuxi.knowledge.utils.kb_utils import parse_minio_url
            from yuxi.storage.minio.client import get_minio_client

            # 解析MinIO URL获取bucket_name和object_name
            bucket_name, object_name = parse_minio_url(file_path)

            # 获取MinIO客户端并下载文件
            minio_client = get_minio_client()
            file_content = await minio_client.adownload_file(bucket_name, object_name)

            # 写入临时文件
            async with aiofiles.open(temp_path, "wb") as f:
                await f.write(file_content)

            logger.debug(f"File downloaded to temp path: {temp_path}")

            # 使用临时文件路径
            actual_file_path = temp_path

        except Exception as e:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            logger.error(f"Failed to download file from MinIO: {e}")
            raise ValueError(f"无法从MinIO下载文件: {e}")
    else:
        # 本地文件
        actual_file_path = file_path

    try:
        file_path_obj = Path(actual_file_path)
        file_ext = file_path_obj.suffix.lower()
        original_filename = file_path_obj.name

        if file_ext == ".pdf":
            # 使用 OCR 处理 PDF
            text = await parse_pdf_async(str(file_path_obj), params=params)
            result = f"{text}"

        elif file_ext in [".txt", ".md"]:
            # 直接读取文本文件
            with open(file_path_obj, encoding="utf-8") as f:
                content = f.read()
            result = f"{content}"

        elif file_ext == ".docx":
            # 优先使用 Docling，失败时回退到 python-docx
            try:
                result = _convert_with_docling(file_path_obj, params=params)
            except Exception as e:
                logger.warning(f"Docling 解析 DOCX 失败，回退到 python-docx: {file_path_obj.name}, {e}")
                result = _convert_docx_with_python_docx(file_path_obj)

        elif file_ext == ".pptx":
            # 使用 Docling 处理 pptx
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".doc":
            # 旧版 .doc 文件仍使用原有解析方式
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader

            loader = UnstructuredWordDocumentLoader(str(file_path_obj))
            docs = loader.load()
            result = "\n".join(doc.page_content for doc in docs).strip()

        elif file_ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"]:
            # 使用 OCR 处理图片
            text = await parse_image_async(str(file_path_obj), params=params)
            result = f"{text}"

        elif file_ext in [".html", ".htm"]:
            # 使用 BeautifulSoup 处理 HTML 文件
            from markdownify import markdownify as md

            with open(file_path_obj, encoding="utf-8") as f:
                content = f.read()
            text = md(content, heading_style="ATX")
            result = f"{text}"

        elif file_ext == ".csv":
            # 处理 CSV 文件
            import pandas as pd

            df = pd.read_csv(file_path_obj)
            # 将每一行数据与表头组合成独立的表格
            markdown_content = ""

            for index, row in df.iterrows():
                # 创建包含表头和当前行的小表格
                row_df = pd.DataFrame([row], columns=df.columns)
                markdown_table = row_df.to_markdown(index=False)
                markdown_content += f"{markdown_table}\n\n"

            result = markdown_content.strip()

        elif file_ext in [".xls", ".xlsx"]:
            # 使用 Docling 处理 Excel 文件
            result = _convert_with_docling(file_path_obj, params=params)

        elif file_ext == ".json":
            # 处理 JSON 文件
            import json

            async with aiofiles.open(file_path_obj, encoding="utf-8") as f:
                content = await f.read()
            data = json.loads(content)
            # 将 JSON 数据格式化为 markdown 代码块
            json_str = json.dumps(data, ensure_ascii=False, indent=2)
            result = f"```json\n{json_str}\n```"

        elif file_ext == ".zip":
            image_bucket, image_prefix = _resolve_image_storage_params(params)
            zip_result = await _process_zip_file(
                str(file_path_obj),
                image_bucket=image_bucket,
                image_prefix=image_prefix,
            )

            # 将处理结果保存到params中供调用方使用
            if params is not None:
                params["_zip_images_info"] = zip_result["images_info"]
                params["_zip_content_hash"] = zip_result["content_hash"]
                params["_zip_image_bucket"] = image_bucket
                params["_zip_image_prefix"] = image_prefix

            result = zip_result["markdown_content"]

        else:
            # 尝试作为文本文件读取
            raise ValueError(f"Unsupported file type: {file_ext}")

    except Exception:
        # 清理临时文件
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as cleanup_e:
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {cleanup_e}")
        raise

    finally:
        # 清理临时文件
        if is_minio_url(file_path) and os.path.exists(actual_file_path):
            try:
                os.unlink(actual_file_path)
                logger.debug(f"Cleaned up temp file: {actual_file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up temp file {actual_file_path}: {e}")

    return result


async def process_url_to_markdown(url: str, params: dict | None = None) -> str:
    """
    Fetch a URL and convert its content to Markdown.

    Args:
        url: The URL to fetch.
        params: Optional parameters (unused, kept for API compatibility).

    Returns:
        The Markdown content of the URL.
    """
    logger.info(f"Fetching URL: {url}")

    try:
        import httpx

        # 使用异步 HTTP 客户端获取页面
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            response.raise_for_status()
            html_content = response.text

        # 使用 readability 提取正文 HTML
        from readability import Document

        doc = Document(html_content)
        body_html = doc.summary()

        # 转换为 Markdown
        markdown_content = md_convert(body_html, heading_style="atx")

        logger.info(f"Successfully converted URL to Markdown: {url}")
        return markdown_content

    except httpx.HTTPError as e:
        logger.error(f"Failed to fetch URL {url}: {e}")
        raise ValueError(f"Failed to fetch URL: {e}")
    except Exception as e:
        logger.error(f"Failed to process URL {url}: {e}")
        raise ValueError(f"Failed to process URL: {e}")
